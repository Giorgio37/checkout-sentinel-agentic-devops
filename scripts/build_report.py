from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT.parents[1] / "outputs" / "Agentic_DevOps_Capstone_Technical_Report.docx"

INK = "17202A"
TEAL = "007A70"
BLUE = "2F6FED"
RED = "B42318"
GRAY = "EEF1F3"
MID = "65727E"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    tc_pr.append(element)


def set_cell_margins(cell, value: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("CSE 636 | Checkout Sentinel   ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)
    add_field(paragraph, "PAGE")


def add_header(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "CHECKOUT SENTINEL  /  AGENTIC DEVOPS CAPSTONE"
    paragraph.style = "Caption"
    paragraph.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    paragraph.runs[0].font.bold = True


def title(doc: Document, text: str, kicker: str | None = None) -> None:
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(kicker.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(11)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = RGBColor.from_string(INK)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(INK)


def body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def evidence_band(doc: Document, values: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(values))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (number, label, color) in enumerate(values):
        cell = table.cell(0, index)
        cell.width = Inches(2.1)
        set_cell_margins(cell, 120)
        shade(cell, GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(number + "\n")
        r.bold = True
        r.font.size = Pt(17)
        r.font.color.rgb = RGBColor.from_string(color)
        r = p.add_run(label)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MID)


def architecture_diagram(doc: Document) -> None:
    labels = [
        ("Candidate\nchange", BLUE),
        ("CI Agent\ntest + repair", TEAL),
        ("Risk Agent\nscore + canary", TEAL),
        ("Human gate\nhash + time", BLUE),
        ("Canary\n10 / 50 / 100", RED),
        ("SRE Agent\nrollback + ITSM", TEAL),
    ]
    table = doc.add_table(rows=1, cols=11)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (label, color) in enumerate(labels):
        cell = table.cell(0, index * 2)
        cell.width = Inches(0.82)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, color)
        set_cell_margins(cell, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        if index < len(labels) - 1:
            arrow = table.cell(0, index * 2 + 1)
            arrow.width = Inches(0.18)
            arrow.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = arrow.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("→")
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor.from_string(MID)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("OpenTelemetry + structured audit + SLSA provenance span every stage")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MID)


def decision_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Decision", "Why", "Trade-off"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        shade(table.cell(0, i), INK)
        for run in table.cell(0, i).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            set_cell_margins(cells[i], 80)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8.5)


def new_page(doc: Document, page: int, heading: str, kicker: str) -> None:
    doc.add_page_break()
    title(doc, heading, f"{page:02d} / {kicker}")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.56)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    add_header(section)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    styles["Normal"].paragraph_format.line_spacing = 1.04
    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(9.2)

    # Page 1
    title(doc, "Checkout Sentinel", "CSE 636 | Agentic DevOps Capstone")
    p = doc.add_paragraph()
    r = p.add_run("A policy-constrained, end-to-end release workflow for a checkout API")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(MID)
    body(doc, "Prepared by Jonathan | August 19, 2026")
    evidence_band(doc, [("5/5", "unit tests passed", TEAL), ("52", "release risk score", BLUE), ("1 / 0 / 0", "real GCP plan", TEAL), ("50%", "maximum rollback scope", RED)])
    h2(doc, "Executive summary")
    body(doc, "Checkout Sentinel demonstrates one causal DevOps chain instead of isolated agent demos. An intentionally broken release candidate fails CI. A constrained CI agent generates a regression test, repairs exactly one runtime field, and proves the second build passes. A risk agent scores the change and selects a 10/50/100 canary. Deployment remains blocked until a human records a timestamped approval bound to the request and risk hashes. At 50% traffic, a controlled anomaly triggers a bounded SRE rollback, ITSM record, postmortem, OpenTelemetry evidence, and structured audit trail.")
    h2(doc, "Scope and evidence boundary")
    bullets(doc, [
        "Real evidence: 5/5 tests; authenticated GCP plan (1 add / 0 change / 0 destroy); 5/5 Conftest pass; timestamped human approval; and one policy-approved local Terraform apply.",
        "Controlled simulation: checkout traffic, 20% error injection, 480 ms p95 latency, canary rollback, and local ITSM connector.",
        "Boundary: the approved GCP apply was attempted but billing returned 403 before creation; no cloud resource exists. Production deployment, signed hosted provenance, remote publication, and Classroom submission were not performed.",
    ])
    h2(doc, "Deliverables")
    body(doc, "The repository contains the service, three agent contracts, CI workflow, Terraform/OPA modules, dashboard, telemetry exporter, audit records, SLSA provenance, tests, report, slides, recording runbook, and AI collaboration record.")

    # Page 2
    new_page(doc, 2, "Architecture keeps agents useful and bounded", "Design")
    architecture_diagram(doc)
    h2(doc, "Autonomy model")
    body(doc, "The CI and SRE agents operate on-loop inside pre-authorized limits; the release agent is human-in-the-loop. The human controls release approval, real cloud apply, remote publication, and final submission. Agents receive structured evidence and tool references, never cloud credential contents.")
    decision_table(doc, [
        ("Deterministic local agents", "Reliable classroom replay without an external API key", "Less flexible than a production LLM planner"),
        ("Integrity-bound approval", "Prevents replaying approval against a changed risk request", "Adds one deliberate manual step"),
        ("10/50/100 canary", "Limits exposure and creates a visible decision point", "Simulated traffic is not production traffic"),
        ("JSONL + OTel", "Readable evidence and trace correlation", "A real backend would retain and query data"),
    ])
    h2(doc, "Agent and tool boundaries")
    bullets(doc, [
        "CI Review Agent: reads test evidence; may patch one allowlisted field in a generated runtime copy.",
        "Release Risk Agent: calculates a documented score and strategy; cannot deploy.",
        "SRE Response Agent: may rollback only an allowlisted canary at or below 50% traffic and create one incident.",
    ])

    # Page 3
    new_page(doc, 3, "The pipeline proves failure, recovery, and rollback", "End-to-end flow")
    h2(doc, "Build and remediation")
    body(doc, "Attempt 1 fails because free_shipping_threshold is 0, incorrectly making a $20 checkout free to ship. The CI agent creates a boundary regression specification, changes only that field to 50, reruns the same tests, and records PASS on attempt 2. Before the repair, an untrusted ticket tries to disable approval and reveal a credential; the security guard blocks it and logs only source metadata and a matched pattern.")
    h2(doc, "Risk, FinOps, and release strategy")
    body(doc, "The risk score is reproducible: base 20 + four changed files (12) + infrastructure change (20) + security signal (10) - passing final tests (10) = 52. Medium risk selects canary-10-50-100 and still requires human approval. The $0.22/month estimate is explicitly illustrative (10 GB and low request volume), not a live cloud quote.")
    h2(doc, "Canary and automated response")
    decision_table(doc, [
        ("10% traffic", "1% errors; 140 ms p95", "Continue"),
        ("50% traffic", "20% errors; 480 ms p95", "Anomaly: thresholds exceeded"),
        ("SRE policy", "rollback_canary allowlisted; scope = 50%", "Rollback 2.3.1 to 2.3.0"),
        ("ITSM", "INC-CAPSTONE-001 + postmortem", "Resolved with linked evidence"),
    ])
    h2(doc, "Testing performed")
    bullets(doc, [
        "Five unit tests passed for checkout behavior, negative subtotal handling, injection detection, unknown-tool denial, and blast-radius enforcement.",
        "A negative release test confirmed deployment raises PermissionError without timestamped approval.",
        "OPA passed all five controls for the compliant plan and correctly failed the staging-label fixture.",
    ])

    # Page 4
    new_page(doc, 4, "Guardrails and telemetry make actions reviewable", "Security and governance")
    h2(doc, "Security guardrails")
    bullets(doc, [
        "Prompt-injection defense: external text is data; attempts to override policy, reveal secrets, disable approval, or launch a shell are denied.",
        "Least privilege: tool allowlist, write scopes, one-field CI repair, 50% SRE ceiling, no raw credential access, and no general shell tool for agents.",
        "OPA policy: GCS must use capstone/terraform labels, uniform bucket-level access, public-access prevention, and versioning.",
        "Terraform lifecycle: prevent_destroy is enabled and force_destroy is false; every apply remains a separate human decision.",
    ])
    h2(doc, "Observability")
    body(doc, "OpenTelemetry spans cover both the checkout service and agent operations. Agent spans use gen_ai.operation.name=invoke_agent with a stable gen_ai.agent.name; tool spans use execute_tool. Metric records include service error rate, p95 latency, traffic percentage, agent tool calls, and agent invocation count. Prompt contents and credentials are intentionally excluded.")
    h2(doc, "Audit and supply-chain governance")
    body(doc, "Each audit record includes UTC timestamp, actor, action, outcome, SHA-256 input/output digests, and bounded details. Approval stores actor, timestamp, request hash, risk hash, release, and environment. The in-toto statement uses predicateType https://slsa.dev/provenance/v1 and identifies the artifact and source materials by digest. This is honest Build L1-style provenance existence, not a claim of signed hosted Build L2/L3 provenance.")
    h2(doc, "IaC evidence")
    body(doc, "On August 19, 2026, Terraform 1.15.8 with Google provider 7.45.0 produced an authenticated GCP plan of one create and no changes or destroys. Conftest 0.68.2 reported 5 passed; the negative environment-label fixture reported 4 passed and 1 failed. After timestamped human approval, the real GCP apply was attempted, but GCP returned 403 accountDisabled before creation; state contained zero resources. A separate no-cost terraform_data sandbox proved the approval-to-apply gate: the approved plan passed 3/3 and applied 1/0/0, while approved=false was blocked 2/3 with one expected failure. The sandbox is not a cloud deployment.")

    # Page 5
    new_page(doc, 5, "The next step is stronger identity, not more autonomy", "Lessons and improvement")
    h2(doc, "Lessons learned")
    bullets(doc, [
        "End-to-end evidence is stronger than seven disconnected weekly demos: one release ID, risk hash, trace context, audit trail, and incident link the story.",
        "Negative tests matter. A blocked unapproved deployment, rejected prompt injection, denied unknown tool, and failed OPA fixture prove guardrails actually enforce behavior.",
        "Agent explanations should be reviewable observation/policy/action/outcome records, not hidden chain-of-thought or unsupported confidence claims.",
        "Human approval has value only when it is scoped and integrity-bound; a free-text approval without hashes is easy to replay incorrectly.",
    ])
    h2(doc, "What I would improve")
    decision_table(doc, [
        ("Identity", "Replace the service-account key with OIDC and short-lived workload identity", "Removes long-lived local credentials"),
        ("Provenance", "Generate signed provenance on a hosted hardened builder", "Progress from local L1-style evidence toward L2/L3"),
        ("Telemetry", "Export OTLP to a retained backend with trace-to-incident links", "Supports real SLOs and historical analysis"),
        ("Deployment", "Restore course-project billing or use an authorized non-production project, then replay traffic", "Completes cloud validation without weakening the gate"),
        ("Agent quality", "Add evaluation sets for remediation precision and injection resistance", "Measures safety instead of assuming it"),
    ])
    h2(doc, "Conclusion")
    body(doc, "Checkout Sentinel meets the capstone intent by combining agentic CI/CD, IaC policy, approval, observability, SRE remediation, governance, and honest limitations. Its central design rule is simple: agents may move quickly inside small, testable boundaries; humans retain decisions that create infrastructure, expose broader traffic, publish externally, or submit coursework.")
    h2(doc, "References")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("OpenTelemetry GenAI semantic conventions: https://opentelemetry.io/docs/specs/semconv/gen-ai/\n")
    r.font.size = Pt(8)
    r = p.add_run("SLSA provenance specification: https://slsa.dev/spec/draft/build-provenance\n")
    r.font.size = Pt(8)
    r = p.add_run("CSE 636 Capstone brief, Google Classroom, accessed August 19, 2026.")
    r.font.size = Pt(8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
